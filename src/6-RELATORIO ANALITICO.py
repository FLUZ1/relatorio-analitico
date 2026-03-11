"""
6-RELATORIO_ANALITICO.py
========================
Gera um relatório PDF analítico inteligente POR PROJETO, com:
  - Resumo executivo
  - Status atual (KPIs globais)
  - Análise de categorias e subcategorias
  - Análise de tendência temporal (backlog)
  - Segmentação Segurança vs Qualidade
  - Destaques críticos e recomendações
"""

import os
import json
import glob
from datetime import datetime
from collections import Counter

import pandas as pd
import numpy as np
import openai

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.lib import colors
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                    Table, TableStyle, HRFlowable, PageBreak,
                                    KeepTogether, Image as RLImage)
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
except ImportError:
    print("\nERRO: pip install reportlab\n")
    import sys; sys.exit(1)

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("\nAVISO: pip install python-docx para gerar arquivos Word (.docx)\n")

try:
    from PyPDF2 import PdfWriter, PdfReader
except ImportError:
    print("\nAVISO: pip install PyPDF2 para anexar relatórios PDF\n")

try:
    from PIL import Image
    from io import BytesIO
except ImportError:
    print("\nAVISO: pip install pillow para compressão de imagens\n")
    Image = None
    BytesIO = None

PASTA_RELATORIOS = "relatorios_analiticos"
PASTA_GRAFICOS   = "graficos_snagr"

def ler_api_key():
    # Chave fornecida diretamente
    return "sk-b61c50e39d2544e68b5ffca511e59796"

API_KEY = ler_api_key()
CLIENT_OPENAI = None

if API_KEY:
    with open("start_log.txt", "a") as f:
        f.write(f"API Key loaded: {API_KEY[:5]}\n")
    try:
        # Configuração para DeepSeek (compatível com OpenAI)
        # Ajuste a base_url se necessário. DeepSeek geralmente usa https://api.deepseek.com/v1
        CLIENT_OPENAI = openai.OpenAI(
            api_key=API_KEY,
            base_url="https://api.deepseek.com",
            timeout=60.0
        )
        with open("start_log.txt", "a") as f:
            f.write("OpenAI Client initialized\n")
    except Exception as e:
        print(f"Erro ao configurar cliente OpenAI: {e}")
else:
    with open("start_log.txt", "a") as f:
        f.write("API Key NOT found\n")

# ─── Cores ──────────────────────────────────────────────────────────────────
COR_PRIM    = colors.HexColor("#01455C")
COR_VERDE   = colors.HexColor("#00e676")
COR_VERM    = colors.HexColor("#ff5252")
COR_CINZA_L = colors.HexColor("#eceff1")
COR_CINZA_M = colors.HexColor("#b0bec5")
BRANCO      = colors.white

# ─── Palavras-chave de segurança ────────────────────────────────────────────
KW_SEG = ['segurança', 'epi', 'epc', 'cipa', 'risco', 'acidente',
          'seguranca', 'sinaliza', 'queda', 'higiene', 'sst']

def classificar_tipo(item):
    txt = f"{item.get('Category','')} {item.get('Defect','')} {item.get('ShortDescrip','')}".lower()
    return 'Segurança' if any(k in txt for k in KW_SEG) else 'Qualidade'

def agrupar_status(s):
    s = str(s or '').strip()
    if s in ['Fixed','Closed','Fechado','Corrigido','Aprovado']: return 'Corrigido'
    if s in ['SignedOff','Verificado']: return 'Verificado'
    if s in ['NotADefect','Recusado','Não Aceito']: return 'Não Aceito'
    return 'Sem Correção'

def parse_data(s):
    if not s or pd.isna(s): return None
    try:
        if 'T' in str(s): return pd.to_datetime(str(s))
        if '/' in str(s): return pd.to_datetime(str(s), format='%d/%m/%Y %H:%M:%S')
    except: pass
    return None

# ─── Integração LLM DeepSeek ────────────────────────────────────────────────
def consultar_llm(contexto):
    """
    Envia o contexto estatístico para o DeepSeek Reasoner e retorna a análise em JSON.
    """
    if not CLIENT_OPENAI:
        return None

    prompt_sistema = """Você é um analista de dados de engenharia civil sênior, especialista em controle de qualidade e segurança.
Sua tarefa é analisar os dados estatísticos fornecidos e gerar os textos para um relatório analítico executivo de alto nível. O projeto está em andamento e os dados são conhecidos, o desafio aqui é engajar e motivar.

Diretrizes Gerais:
1. Autonomia Analítica: Não se limite a descrever os números. Interprete-os, identifique tendências ocultas, correlações e padrões.
2. Contextualização: Este documento tem uma estrutura análitica, com pricipios da observabilidade.No Resumo Executivo, contextualize a situação do projeto de forma abrangente.
3. Proatividade: Na Conclusão, proponha planos de ação concretos, sugira focos de atuação.
4. Ferramentas da Qualidade: Caso identifique padrões relevantes, não sugira ferramentas da qualidade, seja propositivo de maneira genérica e didática, como uma oportunidade de melhoria.
5. Tom de Voz:
    - POSITIVO E CONSTRUTIVO: Evite adjetivos muito negativos ou tom sarcástico. Mesmo diante de dados críticos, foque na oportunidade de melhoria.
    - DIRETIVIDADE: Seja direto e objetivo, sem rodeios.
    - NÃO REPETIÇÃO: Enfatize que o objetivo final é a não repetição dos problemas em serviços futuros (aprendizado organizacional).
6. Precisão: Respeite rigorosamente os números fornecidos no contexto.

Formato de Saída:
- Retorne APENAS um objeto JSON válido.
- Use tags HTML simples (<b>, <i>, <br/>) para formatação.
- NÃO use markdown (```) dentro dos valores do JSON.
- NÃO mencione que este texto foi gerado por uma IA.
- NÃO mencione empreiteira utilize empresa resposável
- Quando da análise houver menção a BENX ou BN, modere os termos e adjetivos.
- Não citar pós obra dando a entender que o projeto está acabando.
- Evite termos em inglês, use português.
Estrutura do JSON:
{
  "resumo": "Resumo executivo contextualizado, apresentando o cenário global do projeto, explique os termos e conceitos utilizados...",
  "segmentacao": "Análise crítica da divisão entre Segurança e Qualidade...",
  "cat_resumo": "Síntese do desempenho por categorias, destacando ofensores e áreas de excelência...",
  "tendencia": "Análise da evolução temporal (backlog), identificando se o ritmo de fechamento acompanha a abertura...",
  "sla": "Análise da eficiência da equipe (MTTR), criticando tempos elevados se houver...",
  "causa_raiz_insights": "Interpretação qualitativa e estratégica dos padrões de causa raiz identificados (NÃO repita os dados, apenas analise e sugira soluções)...",
  "conclusao_final": "Conclusão estratégica abrangente, com recomendações de curto e médio prazo, focando na melhoria contínua e não repetição dos problemas."
}
"""

    prompt_usuario = f"""Dados do Projeto para Análise:
{json.dumps(contexto, indent=2, ensure_ascii=False)}

Gere os textos analíticos conforme as diretrizes."""

    try:
        response = CLIENT_OPENAI.chat.completions.create(
            model="deepseek-chat",
            messages=[
                {"role": "system", "content": prompt_sistema},
                {"role": "user", "content": prompt_usuario}
            ],
            stream=False
        )
        content = response.choices[0].message.content
        # Tentar extrair JSON caso venha com markdown ```json ... ```
        if "```json" in content:
            content = content.split("```json")[1].split("```")[0]
        elif "```" in content:
            content = content.split("```")[1].split("```")[0]
            
        return json.loads(content.strip())
    except Exception as e:
        print(f"Erro na consulta ao LLM: {e}")
        return None

# ─── Estilos ──────────────────────────────────────────────────────────────────
def estilos():
    return {
        'titulo': ParagraphStyle('titulo', fontName='Helvetica-Bold', fontSize=18,
                                  textColor=COR_PRIM, spaceAfter=6*mm, alignment=TA_CENTER),
        'subtitulo': ParagraphStyle('subtitulo', fontName='Helvetica-Bold', fontSize=13,
                                     textColor=COR_PRIM, spaceBefore=6*mm, spaceAfter=3*mm),
        'corpo': ParagraphStyle('corpo', fontName='Helvetica', fontSize=10,
                                 textColor=colors.HexColor("#212121"), leading=15,
                                 alignment=TA_JUSTIFY, spaceAfter=3*mm),
        'destaque': ParagraphStyle('destaque', fontName='Helvetica-Bold', fontSize=10,
                                    textColor=COR_PRIM, leading=14, spaceAfter=2*mm),
        'kpi_valor': ParagraphStyle('kpi_valor', fontName='Helvetica-Bold', fontSize=22,
                                     textColor=COR_PRIM, alignment=TA_CENTER),
        'kpi_label': ParagraphStyle('kpi_label', fontName='Helvetica', fontSize=8,
                                     textColor=COR_CINZA_M, alignment=TA_CENTER),
        'rodape': ParagraphStyle('rodape', fontName='Helvetica', fontSize=7,
                                  textColor=COR_CINZA_M, alignment=TA_CENTER),
    }

# ─── Geração de Textos Analíticos ────────────────────────────────────────────
def gerar_analise(df, nome_projeto):
    """Retorna um dicionário de strings com as análises textuais."""
    total = len(df)
    st = estilos()
    analise = {}

    # 1. Status geral
    df['StatusAg'] = df['FixStatus_Ref'].apply(agrupar_status)
    contagem = df['StatusAg'].value_counts()
    corrigido  = contagem.get('Corrigido', 0)
    verificado = contagem.get('Verificado', 0)
    sem_corr   = contagem.get('Sem Correção', 0)
    nao_aceito = contagem.get('Não Aceito', 0)
    resolvidos = corrigido + verificado
    pct_res    = resolvidos / total * 100 if total else 0
    pct_pend   = sem_corr / total * 100 if total else 0

    analise['resumo'] = (
        f"O projeto <b>{nome_projeto}</b> possui um total de <b>{total} apontamentos</b> registrados. "
        f"Destes, <b>{resolvidos} ({pct_res:.1f}%)</b> foram resolvidos "
        f"(<b>{corrigido}</b> corrigidos + <b>{verificado}</b> verificados), "
        f"enquanto <b>{sem_corr} ({pct_pend:.1f}%)</b> permanecem sem correção."
    )
    if nao_aceito > 0:
        analise['resumo'] += f" Ainda há <b>{nao_aceito}</b> apontamento(s) classificado(s) como 'Não Aceito'."

    if pct_res >= 80:
        analise['resumo'] += (
            f"<br/><br/><b>Avaliação:</b> A taxa de resolução de {pct_res:.0f}% indica um bom desempenho operacional. "
            "A equipe está conseguindo acompanhar o volume de apontamentos gerados."
        )
    elif pct_res >= 50:
        analise['resumo'] += (
            f"<br/><br/><b>Avaliação:</b> A taxa de resolução de {pct_res:.0f}% indica um desempenho moderado. "
            "Recomenda-se intensificar as ações corretivas para reduzir o passivo acumulado."
        )
    else:
        analise['resumo'] += (
            f"<br/><br/><b>Avaliação:</b> A taxa de resolução de {pct_res:.0f}% é crítica. "
            "O volume de pendências está crescendo mais rápido que a capacidade de correção da equipe. "
            "Ação imediata é necessária."
        )

    # 2. Segmentação Segurança vs Qualidade
    df['Tipo'] = df.apply(lambda r: classificar_tipo(r), axis=1)
    tipo_counts = df['Tipo'].value_counts()
    n_seg = tipo_counts.get('Segurança', 0)
    n_qual = tipo_counts.get('Qualidade', 0)
    pct_seg = n_seg / total * 100 if total else 0

    analise['segmentacao'] = (
        f"Do total de apontamentos, <b>{n_qual}</b> são de <b>Qualidade</b> ({n_qual/total*100:.1f}%) e "
        f"<b>{n_seg}</b> são de <b>Segurança</b> ({pct_seg:.1f}%). "
    )
    # Resolução por tipo
    for tipo in ['Qualidade', 'Segurança']:
        df_t = df[df['Tipo'] == tipo]
        if df_t.empty: continue
        st_t = df_t['StatusAg'].value_counts()
        res_t = st_t.get('Corrigido', 0) + st_t.get('Verificado', 0)
        pct_t = res_t / len(df_t) * 100
        analise['segmentacao'] += (
            f"A taxa de resolução de <b>{tipo}</b> é de <b>{pct_t:.1f}%</b> ({res_t}/{len(df_t)}). "
        )

    # 3. Análise por Categoria e Subcategorias (unificado)
    cat_counts = df['Category'].value_counts()
    top_cats = cat_counts.head(8)  # até 8 categorias mais relevantes
    blocos_cat = []

    for cat, n_cat in top_cats.items():
        df_cat = df[df['Category'] == cat]
        res_cat = df_cat['StatusAg'].isin(['Corrigido', 'Verificado']).sum()
        pct_cat = res_cat / n_cat * 100
        pend_cat = n_cat - res_cat

        # Subcategorias dentro desta categoria
        sub_counts = df_cat['Defect'].value_counts()
        linhas_sub = []
        for defect, n_def in sub_counts.items():
            df_def = df_cat[df_cat['Defect'] == defect]
            res_def = df_def['StatusAg'].isin(['Corrigido', 'Verificado']).sum()
            pct_def = res_def / n_def * 100 if n_def else 0
            linhas_sub.append(
                f"&nbsp;&nbsp;&nbsp;&nbsp;• <i>{defect}</i>: {n_def} ocorrências "
                f"({res_def} resolvidos — {pct_def:.0f}%)"
            )

        # Avaliação da categoria
        if pct_cat >= 70:
            avaliacao = "Desempenho satisfatório."
        elif pct_cat >= 40:
            avaliacao = "Desempenho moderado — requer atenção."
        else:
            avaliacao = "<font color='#ff5252'>Desempenho crítico — ação prioritária necessária.</font>"

        bloco = (
            f"<b>{cat}</b> — {n_cat} apontamentos | "
            f"{res_cat} resolvidos ({pct_cat:.0f}%) | {pend_cat} pendentes. "
            f"{avaliacao}<br/>"
            + "<br/>".join(linhas_sub)
        )
        blocos_cat.append(bloco)

    # Categoria mais crítica
    cats_sig = df.groupby('Category').agg(
        total=('SnagID', 'count'),
        resolvidos=('StatusAg', lambda x: x.isin(['Corrigido', 'Verificado']).sum())
    )
    cats_sig['pct'] = cats_sig['resolvidos'] / cats_sig['total'] * 100
    cats_criticas = cats_sig[cats_sig['total'] >= 5].sort_values('pct')
    alerta = ""
    if not cats_criticas.empty:
        pior = cats_criticas.iloc[0]
        alerta = (
            f"<br/><br/><b>Alerta Geral:</b> A categoria com pior desempenho de resolução é "
            f"<b>{cats_criticas.index[0]}</b> com apenas {pior['pct']:.0f}% "
            f"({int(pior['resolvidos'])}/{int(pior['total'])})."
        )

    analise['categorias'] = "<br/><br/>".join(blocos_cat) + alerta

    # ── Texto introdutório e resumo da seção 4 ──
    n_cats = len(top_cats)
    analise['cat_intro'] = (
        f"O gráfico a seguir apresenta a distribuição de apontamentos por "
        f"<b>categoria</b> e suas respectivas <b>subcategorias (defeitos)</b>. "
        f"Para cada categoria, as barras horizontais representam o volume de ocorrências "
        f"por subcategoria, segmentado em <b>resolvidos</b> (verde) e <b>pendentes</b> "
        f"(salmão). O comprimento da barra é proporcional ao item com maior volume de "
        f"ocorrências dentro da categoria. Ao total, {n_cats} categorias mais relevantes "
        f"são apresentadas."
    )

    # Resumo pós-gráfico
    melhor_cat = cats_sig.sort_values('pct', ascending=False).iloc[0] if not cats_sig.empty else None
    txt_resumo = (
        f"<b>Resumo da Análise por Categorias:</b><br/><br/>"
        f"Dos <b>{total} apontamentos</b> registrados, as <b>{n_cats} categorias "
        f"mais frequentes</b> foram analisadas em detalhe. "
    )
    if melhor_cat is not None:
        txt_resumo += (
            f"A categoria com melhor desempenho é <b>{cats_sig.sort_values('pct', ascending=False).index[0]}</b> "
            f"com <b>{melhor_cat['pct']:.0f}%</b> de resolução "
            f"({int(melhor_cat['resolvidos'])}/{int(melhor_cat['total'])}). "
        )
    if not cats_criticas.empty:
        pior = cats_criticas.iloc[0]
        txt_resumo += (
            f"Em contrapartida, a categoria <b>{cats_criticas.index[0]}</b> apresenta "
            f"o pior desempenho com apenas <b>{pior['pct']:.0f}%</b> de resolução "
            f"({int(pior['resolvidos'])}/{int(pior['total'])}), "
            f"demandando atenção prioritária da gestão."
        )
    # Distribuição de avaliações
    n_satisf = int((cats_sig['pct'] >= 70).sum())
    n_moder  = int(((cats_sig['pct'] >= 40) & (cats_sig['pct'] < 70)).sum())
    n_crit   = int((cats_sig['pct'] < 40).sum())
    txt_resumo += (
        f"<br/><br/>Das categorias analisadas, <b>{n_satisf}</b> apresentam desempenho "
        f"satisfatório (≥70%), <b>{n_moder}</b> moderado (40-69%) e "
        f"<b>{n_crit}</b> em nível crítico (<40%)."
    )
    analise['cat_resumo'] = txt_resumo

    # 5. Tendência temporal
    df['Spotted_dt'] = df['Spotted'].apply(parse_data)
    df_com_data = df.dropna(subset=['Spotted_dt'])
    if not df_com_data.empty:
        por_mes = df_com_data.set_index('Spotted_dt').resample('ME').size()
        if len(por_mes) >= 2:
            ultimo_mes = por_mes.iloc[-1]
            penultimo  = por_mes.iloc[-2]
            media      = por_mes.mean()
            nome_ultimo    = por_mes.index[-1].strftime('%B/%Y')
            nome_penultimo = por_mes.index[-2].strftime('%B/%Y')

            # ── Mês atual ──
            tendencia = (
                f"<b>Mês atual — {nome_ultimo}:</b> Foram registrados "
                f"<b>{ultimo_mes} apontamentos</b> neste período. "
            )

            # ── Comparação com mês anterior ──
            if penultimo > 0:
                var_pct = ((ultimo_mes - penultimo) / penultimo) * 100
                if var_pct > 0:
                    tendencia += (
                        f"Em comparação com o mês anterior ({nome_penultimo}), que teve "
                        f"<b>{penultimo} apontamentos</b>, houve um "
                        f"<font color='#ff5252'><b>aumento de {var_pct:.0f}%</b></font>. "
                        f"O crescimento no volume de apontamentos indica possível "
                        f"aceleração na identificação de não-conformidades ou "
                        f"aumento real de defeitos."
                    )
                elif var_pct < 0:
                    tendencia += (
                        f"Em comparação com o mês anterior ({nome_penultimo}), que teve "
                        f"<b>{penultimo} apontamentos</b>, houve uma "
                        f"<font color='#00c853'><b>redução de {abs(var_pct):.0f}%</b></font>. "
                        f"A queda pode indicar melhoria nos processos construtivos "
                        f"ou redução no ritmo de inspeções."
                    )
                else:
                    tendencia += (
                        f"O volume manteve-se estável em relação ao mês anterior "
                        f"({nome_penultimo}: {penultimo} apontamentos)."
                    )
            else:
                tendencia += (
                    f"O mês anterior ({nome_penultimo}) não possuía registros "
                    f"para comparação direta."
                )

            # ── Tendência geral vs média ──
            if ultimo_mes > media * 1.3:
                tendencia += (
                    f"<br/><br/>O mês atual está <b>acima da média mensal de "
                    f"{media:.0f} apontamentos</b>. A tendência é crescente — "
                    f"o volume de problemas está aumentando."
                )
            elif ultimo_mes < media * 0.7:
                tendencia += (
                    f"<br/><br/>O mês atual está <b>abaixo da média mensal de "
                    f"{media:.0f} apontamentos</b>. A tendência é de queda — "
                    f"o ritmo de novos problemas está desacelerando."
                )
            else:
                tendencia += (
                    f"<br/><br/>O mês atual está <b>próximo da média mensal de "
                    f"{media:.0f} apontamentos</b>. A tendência é estável."
                )

            # Período com mais apontamentos
            pico = por_mes.idxmax()
            tendencia += (
                f"<br/><br/>O pico de apontamentos ocorreu em <b>{pico.strftime('%B/%Y')}</b> "
                f"com <b>{por_mes.max()}</b> registros."
            )

            # ── Histórico acumulado e volume de correção ──
            total_meses = len(por_mes)
            primeiro_mes = por_mes.index[0].strftime('%B/%Y')
            total_abertos  = int(por_mes.sum())
            # Correções ao longo do tempo
            df_res = df[df['StatusAg'].isin(['Corrigido', 'Verificado'])].copy()
            df_res['LastChanged_dt2'] = df_res['LastChanged'].apply(parse_data)
            df_res_data = df_res.dropna(subset=['LastChanged_dt2'])
            if not df_res_data.empty:
                corr_mes = df_res_data.set_index('LastChanged_dt2').resample('ME').size()
                media_corr = corr_mes.mean()
                total_corrigidos = int(corr_mes.sum())
                tendencia += (
                    f"<br/><br/><b>Histórico acumulado:</b> Desde <b>{primeiro_mes}</b>, "
                    f"ao longo de <b>{total_meses} meses</b>, foram registrados "
                    f"<b>{total_abertos} apontamentos</b>. Neste mesmo período, "
                    f"<b>{total_corrigidos} itens foram corrigidos</b>, com uma média "
                    f"de <b>{media_corr:.0f} correções/mês</b>. "
                )
                if media_corr >= media:
                    tendencia += (
                        f"O ritmo de correção (<b>{media_corr:.0f}/mês</b>) é "
                        f"igual ou superior ao de abertura (<b>{media:.0f}/mês</b>), "
                        f"indicando que a equipe está conseguindo acompanhar a demanda."
                    )
                else:
                    deficit = media - media_corr
                    tendencia += (
                        f"O ritmo de correção (<b>{media_corr:.0f}/mês</b>) está "
                        f"abaixo do de abertura (<b>{media:.0f}/mês</b>), "
                        f"gerando um acúmulo estimado de <b>{deficit:.0f} itens/mês</b> "
                        f"no passivo pendente. Ação de reforço é recomendada."
                    )
            else:
                tendencia += (
                    f"<br/><br/><b>Histórico acumulado:</b> Desde <b>{primeiro_mes}</b>, "
                    f"ao longo de <b>{total_meses} meses</b>, foram registrados "
                    f"<b>{total_abertos} apontamentos</b>. Dados de correção "
                    f"insuficientes para análise do volume de resolução ao longo do tempo."
                )

            analise['tendencia'] = tendencia
        else:
            analise['tendencia'] = "Dados temporais insuficientes para análise de tendência (menos de 2 meses)."
    else:
        analise['tendencia'] = "Sem datas de abertura disponíveis para análise temporal."

    # 8. SLA / Tempo de Resposta (MTTR)
    df['LastChanged_dt'] = df['LastChanged'].apply(parse_data)
    df_sla = df.dropna(subset=['Spotted_dt', 'LastChanged_dt']).copy()
    if not df_sla.empty:
        df_sla['SLA_dias'] = (df_sla['LastChanged_dt'] - df_sla['Spotted_dt']).dt.days

        # MTTR global
        mttr_global = df_sla['SLA_dias'].mean()
        mediana_global = df_sla['SLA_dias'].median()
        maximo = df_sla['SLA_dias'].max()

        txt_sla = (
            "A análise de <b>SLA (Service Level Agreement)</b> e <b>MTTR (Mean Time To Repair)</b> "
            "é fundamental para medir a eficiência da equipe na resolução dos problemas identificados. "
            "O SLA define o compromisso de tempo para a tratativa, enquanto o MTTR reflete a média real "
            "de dias que um apontamento leva para ser concluído, desde sua abertura até a última interação.<br/><br/>"
            f"O tempo médio entre abertura e última interação (MTTR) é de "
            f"<b>{mttr_global:.0f} dias</b> (mediana: {mediana_global:.0f} dias, máximo: {maximo} dias)."
        )

        # MTTR por categoria (top 5) - Preparar dados para tabela
        sla_cat_data = []
        sla_cat = df_sla.groupby('Category')['SLA_dias'].agg(['mean', 'median', 'count'])
        sla_cat = sla_cat[sla_cat['count'] >= 3].sort_values('mean', ascending=False).head(5)
        if not sla_cat.empty:
            for cat, row in sla_cat.iterrows():
                alerta_sla = ""
                if row['mean'] > 60:
                    alerta_sla = "Crítico"
                elif row['mean'] > 30:
                    alerta_sla = "Atenção"
                sla_cat_data.append({
                    'Categoria': cat,
                    'Média': f"{row['mean']:.0f} dias",
                    'Mediana': f"{row['median']:.0f} dias",
                    'Qtd': int(row['count']),
                    'Status': alerta_sla
                })

        # MTTR por grupo/empreiteira (top 5) - Preparar dados para tabela
        sla_grp_data = []
        sla_grp = df_sla.groupby('Groupname')['SLA_dias'].agg(['mean', 'median', 'count'])
        sla_grp = sla_grp[sla_grp['count'] >= 3].sort_values('mean', ascending=False).head(5)
        if not sla_grp.empty:
            for grp, row in sla_grp.iterrows():
                alerta_sla = ""
                if row['mean'] > 60:
                    alerta_sla = "Crítico"
                elif row['mean'] > 30:
                    alerta_sla = "Atenção"
                sla_grp_data.append({
                    'Grupo': grp,
                    'Média': f"{row['mean']:.0f} dias",
                    'Mediana': f"{row['median']:.0f} dias",
                    'Qtd': int(row['count']),
                    'Status': alerta_sla
                })

        # Itens abertos há mais tempo (Top 20 para tabela)
        df_pendentes = df_sla[df_sla['StatusAg'] == 'Sem Correção'].copy()
        top_antigos_data = []
        if not df_pendentes.empty:
            df_pendentes['Dias_aberto'] = (pd.Timestamp.now() - df_pendentes['Spotted_dt']).dt.days
            top_antigos = df_pendentes.nlargest(20, 'Dias_aberto')
            
            for _, row in top_antigos.iterrows():
                top_antigos_data.append({
                    'ID': row.get('SnagID', ''),
                    'Dias': row['Dias_aberto'],
                    'Categoria': row.get('Category', ''),
                    'Defeito': row.get('Defect', ''),
                    'Local': row.get('Location', '')
                })

        analise['sla'] = txt_sla
        analise['sla_cat_data'] = sla_cat_data
        analise['sla_grp_data'] = sla_grp_data
        analise['top_antigos'] = top_antigos_data
    else:
        analise['sla'] = "Dados insuficientes para cálculo de SLA (datas ausentes)."
        analise['sla_cat_data'] = []
        analise['sla_grp_data'] = []
        analise['top_antigos'] = []


    # 9. Análise de Causa Raiz (Heurística Nativa)
    causas = []
    
    analise['causa_raiz_intro'] = (
        "A <b>Análise de Causa Raiz</b> é uma forma prática de olhar para os dados e entender "
        "o que está se repetindo com mais frequência — ou seja, onde os problemas tendem a voltar."
        "<br/><br/>"
        "Nesta seção, a ideia não é só apontar “qual foi o defeito”, mas enxergar padrões de "
        "reincidência entre a empresa responsável, o tipo de apontamento e os locais mais afetados. "
        "Com isso, fica mais fácil direcionar ações que ataquem a origem do problema e reduzam a "
        "repetição ao longo do tempo."
    )

    # 9a. Concentração: Empreiteira × Defeito
    if 'Groupname' in df.columns and 'Defect' in df.columns:
        cross_gd = df.groupby(['Groupname', 'Defect']).size().reset_index(name='qtd')
        cross_gd = cross_gd[cross_gd['qtd'] >= 3].sort_values('qtd', ascending=False).head(5)
        if not cross_gd.empty:
            linhas_gd = []
            for _, row in cross_gd.iterrows():
                df_sub = df[(df['Groupname'] == row['Groupname']) & (df['Defect'] == row['Defect'])]
                res_sub = df_sub['StatusAg'].isin(['Corrigido', 'Verificado']).sum()
                pct_sub = res_sub / row['qtd'] * 100
                linhas_gd.append(
                    f"&nbsp;&nbsp;&nbsp;&nbsp;• <b>{row['Groupname']}</b> × <i>{row['Defect']}</i>: "
                    f"{row['qtd']} ocorrências ({pct_sub:.0f}% resolvidos)"
                )
            causas.append(
                "<b>Concentração Empreiteira × Defeito:</b> Algumas combinações "
                "de empreiteira e tipo de defeito se repetem sistematicamente, "
                "indicando possível falha de método ou capacitação:<br/>"
                + "<br/>".join(linhas_gd)
            )

    # 9b. Concentração: Localização × Categoria
    loc_col = 'Location' if 'Location' in df.columns else 'DrwgTitle'
    if loc_col in df.columns:
        cross_lc = df.groupby([loc_col, 'Category']).size().reset_index(name='qtd')
        cross_lc = cross_lc[cross_lc['qtd'] >= 3].sort_values('qtd', ascending=False).head(5)
        if not cross_lc.empty:
            linhas_lc = []
            for _, row in cross_lc.iterrows():
                linhas_lc.append(
                    f"&nbsp;&nbsp;&nbsp;&nbsp;• <b>{row[loc_col]}</b> × <i>{row['Category']}</i>: "
                    f"{row['qtd']} ocorrências"
                )
            causas.append(
                "<b>Concentração Localização × Categoria:</b> Determinadas áreas "
                "acumulam defeitos de mesma natureza, sugerindo causa sistêmica "
                "(material, projeto ou mão de obra local):<br/>"
                + "<br/>".join(linhas_lc)
            )

    # 9c. Defeitos crônicos (abertos > 90 dias, sem correção)
    if 'Spotted_dt' in df.columns:
        df_cronicos = df[(df['StatusAg'] == 'Sem Correção')].copy()
        df_cronicos['Spotted_dt2'] = df_cronicos['Spotted'].apply(parse_data)
        df_cronicos = df_cronicos.dropna(subset=['Spotted_dt2'])
        if not df_cronicos.empty:
            df_cronicos['Dias'] = (pd.Timestamp.now() - df_cronicos['Spotted_dt2']).dt.days
            cronicos = df_cronicos[df_cronicos['Dias'] > 90]
            if len(cronicos) > 0:
                # Agrupar por categoria
                cron_cat = cronicos['Category'].value_counts().head(5)
                linhas_cron = [
                    f"&nbsp;&nbsp;&nbsp;&nbsp;• <b>{cat}</b>: {n} itens crônicos (> 90 dias sem correção)"
                    for cat, n in cron_cat.items()
                ]
                causas.append(
                    f"<b>Defeitos Crônicos ({len(cronicos)} itens):</b> "
                    f"Apontamentos abertos há mais de 90 dias sem tratativa indicam "
                    "problemas estruturais de gestão — falta de responsável, "
                    "indefinição técnica ou abandono do acompanhamento:<br/>"
                    + "<br/>".join(linhas_cron)
                )

    # 9d. Reincidência temporal (mesma subcategoria aparece em >3 meses distintos)
    if 'Spotted_dt' in df.columns:
        df_temp = df.dropna(subset=['Spotted_dt']).copy()
        if not df_temp.empty:
            df_temp['Mes'] = df_temp['Spotted_dt'].dt.to_period('M')
            reincid = df_temp.groupby('Defect')['Mes'].nunique()
            reincid = reincid[reincid >= 3].sort_values(ascending=False).head(5)
            if not reincid.empty:
                linhas_r = [
                    f"&nbsp;&nbsp;&nbsp;&nbsp;• <b>{defect}</b>: aparece em <b>{n_meses}</b> meses distintos"
                    for defect, n_meses in reincid.items()
                ]
                causas.append(
                    "<b>Reincidência Temporal:</b> Defeitos que se repetem ao longo "
                    "de vários meses indicam que a causa raiz não foi tratada — "
                    "apenas o sintoma está sendo corrigido pontualmente:<br/>"
                    + "<br/>".join(linhas_r)
                )

    if causas:
        analise['causa_raiz'] = "<br/><br/>".join(causas)
    else:
        analise['causa_raiz'] = (
            "Não foram identificados padrões de concentração significativa nos dados. "
            "Base de dados atual insuficiente para análise heurística de causa raiz."
        )
    
    # Seção de conclusão de causa raiz removida conforme solicitado
    analise['causa_raiz_conclusao'] = ""

    # ─── Integração LLM (Substituição dos Textos) ───
    if CLIENT_OPENAI:
        try:
            print(f"  [{nome_projeto}] Consultando IA...")
            # Preparar contexto
            kpi_context = {
                "projeto": nome_projeto,
                "total_apontamentos": int(total),
                "resolvidos": int(resolvidos),
                "pendentes": int(sem_corr),
                "taxa_resolucao": round(pct_res, 1),
                "seguranca": int(n_seg),
                "qualidade": int(n_qual),
            }
            
            # Categorias principais
            cats_context = []
            # cats_sig é calculado anteriormente na função
            if 'cats_sig' in locals() and not cats_sig.empty:
                top_c = cats_sig.sort_values('total', ascending=False).head(10)
                for c, row in top_c.iterrows():
                    cats_context.append({
                        "categoria": str(c),
                        "total": int(row['total']),
                        "resolvidos": int(row['resolvidos']),
                        "pct": round(row['pct'], 1)
                    })
            kpi_context['categorias_top'] = cats_context
            
            # Tendência (se disponível)
            if 'por_mes' in locals() and len(por_mes) >= 2:
                kpi_context['tendencia_mensal'] = {
                    "ultimo_mes": int(por_mes.iloc[-1]),
                    "mes_anterior": int(por_mes.iloc[-2]),
                    "media_mensal": round(por_mes.mean(), 1)
                }
                
            # SLA (se disponível)
            if 'mttr_global' in locals():
                kpi_context['sla_mttr_dias'] = round(mttr_global, 1)
                
            # Causas Raiz (Resumo das encontradas nativamente)
            if causas:
                kpi_context['padroes_causa_raiz'] = causas

            # Consultar LLM
            analise_llm = consultar_llm(kpi_context)
            
            if analise_llm:
                # Atualizar chaves se existirem no retorno
                for k, v in analise_llm.items():
                    if v:
                        if k == 'causa_raiz_insights' and 'causa_raiz' in analise:
                            # Concatenar insights da IA após os dados nativos
                            analise['causa_raiz'] += f"<br/><br/><b>Análise e Recomendações:</b><br/>{v}"
                        elif k != 'causa_raiz_insights': # Evitar criar chave duplicada ou errada
                            analise[k] = v

        except Exception as e:
            print(f"Erro ao processar LLM: {e}")

    return analise, {
        'total': total, 'resolvidos': resolvidos, 'sem_corr': sem_corr,
        'pct_res': pct_res, 'n_seg': n_seg, 'n_qual': n_qual
    }

# ─── Anexar PDF ────────────────────────────────────────────────────────────
def encontrar_ultimo_relatorio_snagr(nome_proj):
    """Encontra o PDF mais recente na pasta relatorios_SnagR para um projeto."""
    pasta_proj = os.path.join("relatorios_SnagR", nome_proj)
    if not os.path.exists(pasta_proj):
        return None
    
    pastas_data = sorted(glob.glob(os.path.join(pasta_proj, "20*")), reverse=True)
    if not pastas_data:
        return None
    
    arquivo_pdf = os.path.join(pastas_data[0], f"{nome_proj}_Relatorio.pdf")
    return arquivo_pdf if os.path.exists(arquivo_pdf) else None

# ─── Gerar Conclusão Final ─────────────────────────────────────────────────
def gerar_conclusao(analise, kpis, nome_projeto):
    # Se o LLM gerou uma conclusão final robusta, use-a.
    if analise.get('conclusao_final'):
        return analise['conclusao_final']

    total = kpis['total']
    pct = kpis['pct_res']
    pend = kpis['sem_corr']

    # Tom geral
    if pct >= 80:
        abertura = (
            f"O projeto <b>{nome_projeto}</b> apresenta um cenário positivo de gestão da qualidade. "
            f"Com <b>{pct:.0f}%</b> dos apontamentos resolvidos, a equipe demonstra comprometimento "
            "e capacidade de resposta às demandas identificadas em campo."
        )
        perspectiva = (
            "A perspectiva é favorável para a continuidade da operação. "
            "Recomenda-se manter o ritmo de tratativas e focar na resolução dos itens "
            "crônicos remanescentes para atingir a excelência operacional."
        )
    elif pct >= 50:
        abertura = (
            f"O projeto <b>{nome_projeto}</b> encontra-se em um estágio intermediário de gestão. "
            f"A taxa de resolução de <b>{pct:.0f}%</b> indica que há esforço de tratativa, "
            f"porém o volume de <b>{pend} itens pendentes</b> exige atenção redobrada."
        )
        perspectiva = (
            "É importante intensificar os planos de ação nas categorias mais críticas "
            "e estabelecer metas semanais de resolução. A aplicação de causa raiz "
            "nos defeitos reincidentes pode reduzir significativamente o passivo."
        )
    else:
        abertura = (
            f"O projeto <b>{nome_projeto}</b> apresenta um cenário que demanda ação imediata. "
            f"Com apenas <b>{pct:.0f}%</b> de resolução e <b>{pend} apontamentos pendentes</b>, "
            "o acúmulo de não-conformidades representa risco à qualidade final da entrega "
            "e à segurança do canteiro."
        )
        perspectiva = (
            "Recomenda-se fortemente a elaboração de um planto de ação "
            "para correção dos itens críticos, definição de responsáveis e "
            "prazos realistas. O acompanhamento mais focado é indispensável nesta fase, para correção dos itens apontados e principalmente e evitar a reincidência em outros pavimentos."
        )

    conclusao = (
        f"{abertura}"
        f"<br/><br/>Ao longo deste relatório, foram analisados <b>{total} apontamentos</b> "
        "sob as perspectivas de segmentação (Qualidade vs Segurança), desempenho por categoria "
        "e subcategoria, tendência temporal, SLA de resposta e padrões de causa raiz. "
        "Os dados permitem uma visão objetiva do estado atual e orientam as próximas decisões."
        f"<br/><br/>{perspectiva}"
        "<br/><br/><i>Este relatório foi gerado automaticamente a partir dos dados registrados "
        "na plataforma SnagR, utilizando análise estatística nativa. As conclusões "
        "e recomendações são baseadas exclusivamente nos padrões identificados nos dados.</i>"
    )
    return conclusao

# ─── Seção 4 — Mini Gráficos de Barras por Subcategoria ────────────────────
def montar_secao_categorias(df, largura):
    """Gera flowables com mini gráficos de barras horizontais por subcategoria."""
    from reportlab.graphics.shapes import Drawing, Rect, String as GStr

    flowables = []

    LABEL_W = 62 * mm      # coluna do rótulo
    VAL_W   = 22 * mm      # coluna do valor
    BAR_W   = largura - LABEL_W - VAL_W  # coluna da barra
    ROW_H   = 7  * mm      # altura de cada linha de subcategoria
    BAR_H   = 4  * mm      # espessura da barra

    # ── Legenda ──────────────────────────────────────────────────────────────
    leg = Drawing(largura, 8 * mm)
    leg.add(Rect(0,       1*mm, 8*mm, 4*mm, fillColor=COR_VERDE,                  strokeColor=None))
    leg.add(GStr(10*mm,   2*mm, "Resolvidos",  fontSize=7, fillColor=colors.HexColor("#555555")))
    leg.add(Rect(50*mm,   1*mm, 8*mm, 4*mm, fillColor=colors.HexColor("#ffccbc"), strokeColor=None))
    leg.add(GStr(60*mm,   2*mm, "Pendentes",   fontSize=7, fillColor=colors.HexColor("#555555")))
    leg.add(Rect(100*mm,  1*mm, 8*mm, 4*mm, fillColor=COR_CINZA_L,               strokeColor=COR_CINZA_M, strokeWidth=0.4))
    leg.add(GStr(110*mm,  2*mm, "Sem dados na faixa", fontSize=7, fillColor=colors.HexColor("#888888")))
    flowables.append(leg)
    flowables.append(Spacer(1, 3 * mm))

    cat_counts = df["Category"].value_counts()
    top_cats   = cat_counts.head(8)

    for cat, n_cat in top_cats.items():
        df_cat  = df[df["Category"] == cat].copy()
        res_cat = int(df_cat["StatusAg"].isin(["Corrigido", "Verificado"]).sum())
        pct_cat = res_cat / n_cat * 100 if n_cat else 0

        if pct_cat >= 70:
            cor_badge = COR_VERDE
            avaliacao = "Satisfatório"
        elif pct_cat >= 40:
            cor_badge = colors.HexColor("#ef6c00")
            avaliacao = "Moderado"
        else:
            cor_badge = COR_VERM
            avaliacao = "Crítico"

        # ── Cabeçalho da Categoria ────────────────────────────────────────────
        st_h = ParagraphStyle("ch", fontName="Helvetica-Bold", fontSize=9, textColor=BRANCO)
        st_p = ParagraphStyle("cp", fontName="Helvetica",      fontSize=8, textColor=BRANCO)

        hdr = Table(
            [[Paragraph(str(cat), st_h),
              Paragraph(f"{n_cat} itens | {pct_cat:.0f}% resolvidos | {avaliacao}", st_p)]],
            colWidths=[largura * 0.55, largura * 0.45]
        )
        hdr.setStyle(TableStyle([
            ("BACKGROUND",    (0, 0), (0, 0),  COR_PRIM),
            ("BACKGROUND",    (1, 0), (1, 0),  cor_badge),
            ("VALIGN",        (0, 0), (-1, -1), "MIDDLE"),
            ("LEFTPADDING",   (0, 0), (-1, -1), 8),
            ("RIGHTPADDING",  (0, 0), (-1, -1), 8),
            ("TOPPADDING",    (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ("ALIGN",         (1, 0), (1,  0),  "RIGHT"),
        ]))

        # ── Barras por Subcategoria ───────────────────────────────────────────
        sub_counts = df_cat["Defect"].value_counts()
        if sub_counts.empty:
            flowables.append(KeepTogether([hdr, Spacer(1, 3 * mm)]))
            continue

        max_val  = float(sub_counts.max())
        n_subs   = len(sub_counts)
        total_h  = n_subs * ROW_H + 3 * mm

        d = Drawing(largura, total_h)

        for i, (defect, n_def) in enumerate(sub_counts.items()):
            y_base  = total_h - (i + 1) * ROW_H + 1 * mm
            bar_y   = y_base + 1 * mm
            prop    = n_def / max_val
            fill_w  = BAR_W * prop

            df_def  = df_cat[df_cat["Defect"] == defect]
            res_def = int(df_def["StatusAg"].isin(["Corrigido", "Verificado"]).sum())
            pct_d   = res_def / n_def if n_def else 0

            # Track (fundo cinza)
            d.add(Rect(LABEL_W, bar_y, BAR_W, BAR_H,
                       fillColor=COR_CINZA_L, strokeColor=COR_CINZA_M, strokeWidth=0.3))

            # Parte pendente (salmão)
            if fill_w > 0:
                d.add(Rect(LABEL_W, bar_y, fill_w, BAR_H,
                           fillColor=colors.HexColor("#ffccbc"), strokeColor=None))

            # Parte resolvida (verde)
            res_w = fill_w * pct_d
            if res_w > 0:
                d.add(Rect(LABEL_W, bar_y, res_w, BAR_H,
                           fillColor=COR_VERDE, strokeColor=None))

            # Rótulo (truncado em 42 chars)
            lbl = str(defect)[:42] if defect else "—"
            d.add(GStr(2 * mm, y_base + 1.5 * mm, lbl,
                       fontSize=7, fillColor=colors.HexColor("#333333")))

            # Valor numérico
            d.add(GStr(LABEL_W + BAR_W + 2 * mm, y_base + 1.5 * mm,
                       f"{n_def}  ({pct_d * 100:.0f}%)",
                       fontSize=7, fillColor=colors.HexColor("#444444")))

        # Manter cabeçalho + gráfico juntos na mesma página
        flowables.append(KeepTogether([hdr, d, Spacer(1, 4 * mm)]))

    return flowables


def montar_grafico_tendencia(df, largura):
    from reportlab.graphics.shapes import Drawing, Rect, String as GStr, Line as GLine, PolyLine, Circle
    from reportlab.platypus import Table, Spacer
    import math

    df2 = df.copy()
    if 'Spotted_dt' not in df2.columns:
        df2['Spotted_dt'] = df2['Spotted'].apply(parse_data)
    else:
        df2['Spotted_dt'] = df2['Spotted_dt'].apply(lambda x: x if not pd.isna(x) else None)

    df_open = df2.dropna(subset=['Spotted_dt'])
    if df_open.empty:
        return []

    abertos = df_open.set_index('Spotted_dt').resample('ME').size()

    if 'StatusAg' not in df2.columns:
        if 'FixStatus_Ref' in df2.columns:
            df2['StatusAg'] = df2['FixStatus_Ref'].apply(agrupar_status)
        else:
            df2['StatusAg'] = df2.apply(lambda r: agrupar_status(r.get('FixStatus') or r.get('Status', 'Open')), axis=1)

    df_res = df2[df2['StatusAg'].isin(['Corrigido', 'Verificado'])].copy()
    df_res['LastChanged_dt2'] = df_res['LastChanged'].apply(parse_data)
    df_res = df_res.dropna(subset=['LastChanged_dt2'])
    resolvidos = df_res.set_index('LastChanged_dt2').resample('ME').size() if not df_res.empty else pd.Series(dtype=int)

    idx = abertos.index.union(resolvidos.index).sort_values()
    if len(idx) < 2:
        return []

    idx = idx[-12:]

    abertos_v = [int(abertos.get(m, 0)) for m in idx]
    resolvidos_v = [int(resolvidos.get(m, 0)) for m in idx]

    backlog_v = []
    backlog = 0
    for a, r in zip(abertos_v, resolvidos_v):
        backlog = backlog + a - r
        backlog_v.append(max(0, int(backlog)))

    max_y = max(abertos_v + resolvidos_v + backlog_v + [1])
    max_y = int(math.ceil(max_y / 10.0) * 10) if max_y >= 10 else max_y

    w = float(largura)
    h = 55 * mm
    m_left = 18 * mm
    m_right = 4 * mm
    m_bottom = 10 * mm
    m_top = 10 * mm

    x0 = m_left
    y0 = m_bottom
    x1 = w - m_right
    y1 = h - m_top

    d = Drawing(w, h)
    d.add(Rect(0, 0, w, h, fillColor=colors.white, strokeColor=COR_CINZA_L, strokeWidth=0.6))
    d.add(Rect(x0, y0, x1 - x0, y1 - y0, fillColor=colors.white, strokeColor=COR_CINZA_L, strokeWidth=0.6))

    ticks = 4
    for i in range(ticks + 1):
        v = int(round((max_y * i) / ticks))
        y = y0 + ((y1 - y0) * i / ticks)
        d.add(GLine(x0, y, x1, y, strokeColor=COR_CINZA_L, strokeWidth=0.5))
        d.add(GStr(2, y - 2, str(v), fontSize=7, fillColor=COR_CINZA_M))

    n = len(idx)
    denom = (n - 1) if n > 1 else 1

    def xy(i, v):
        x = x0 + (x1 - x0) * (i / denom)
        y = y0 + (y1 - y0) * (v / max_y) if max_y else y0
        return x, y

    cor_abertos = COR_PRIM
    cor_res = COR_VERDE
    cor_backlog = colors.HexColor("#ef6c00")

    def add_series(vals, cor):
        pts = []
        for i, v in enumerate(vals):
            x, y = xy(i, v)
            pts.extend([x, y])
        d.add(PolyLine(pts, strokeColor=cor, strokeWidth=1.6))
        for i, v in enumerate(vals):
            x, y = xy(i, v)
            d.add(Circle(x, y, 1.4, fillColor=cor, strokeColor=cor))

    add_series(abertos_v, cor_abertos)
    add_series(resolvidos_v, cor_res)
    add_series(backlog_v, cor_backlog)

    leg_y = h - 6 * mm
    leg_x = x0
    leg_items = [
        (cor_abertos, "Abertos"),
        (cor_res, "Resolvidos"),
        (cor_backlog, "Backlog"),
    ]
    for cor, lbl in leg_items:
        d.add(GLine(leg_x, leg_y, leg_x + 10 * mm, leg_y, strokeColor=cor, strokeWidth=2.0))
        d.add(GStr(leg_x + 12 * mm, leg_y - 2, lbl, fontSize=8, fillColor=colors.HexColor("#444444")))
        leg_x += 32 * mm

    step = max(1, int(math.ceil(n / 6.0)))
    for i, m in enumerate(idx):
        if i % step != 0 and i != n - 1:
            continue
        x, _ = xy(i, 0)
        d.add(GStr(x - 8, 2, m.strftime('%m/%y'), fontSize=7, fillColor=colors.HexColor("#666666")))

    linhas = min(6, n)
    idx_tbl = idx[-linhas:]
    data_tbl = [["Mês", "Abertos", "Resolvidos", "Backlog"]]
    start = n - linhas
    for j, m in enumerate(idx_tbl):
        i = start + j
        data_tbl.append([m.strftime('%m/%Y'), str(abertos_v[i]), str(resolvidos_v[i]), str(backlog_v[i])])

    t = Table(data_tbl, colWidths=[largura * 0.22, largura * 0.18, largura * 0.20, largura * 0.20])
    t.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), COR_PRIM),
        ('TEXTCOLOR', (0, 0), (-1, 0), BRANCO),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 8),
        ('ALIGN', (1, 0), (-1, -1), 'CENTER'),
        ('ALIGN', (0, 0), (0, -1), 'LEFT'),
        ('GRID', (0, 0), (-1, -1), 0.4, COR_CINZA_M),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [BRANCO, COR_CINZA_L]),
        ('FONTSIZE', (0, 1), (-1, -1), 8),
        ('TOPPADDING', (0, 0), (-1, -1), 3),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
    ]))

    return [Spacer(1, 2 * mm), d, Spacer(1, 2 * mm), t, Spacer(1, 2 * mm)]


# ─── Montar PDF ────────────────────────────────────────────────────────────
def _on_cover(nome_projeto, data_impressao):
    """Callback para a página de capa — usa imagem modelo como fundo."""
    def fn(canvas, doc):
        w, h = A4
        canvas.saveState()

        # ── Caminho da imagem modelo de capa ──
        script_dir = os.path.dirname(os.path.abspath(__file__))
        img_capa = os.path.join(script_dir, "Imagem Modelo",
                                "modelo capa_relatório análitico.png")

        if os.path.exists(img_capa):
            # Imagem modelo cobre toda a página A4
            # Otimização: compressão JPEG para reduzir tamanho do PDF
            if Image is not None and BytesIO is not None:
                try:
                    # Abrir imagem com Pillow
                    pil_img = Image.open(img_capa)
                    # Converter para RGB (caso PNG com canal alpha)
                    if pil_img.mode in ('RGBA', 'LA', 'P'):
                        background = Image.new('RGB', pil_img.size, (255, 255, 255))
                        background.paste(pil_img, mask=pil_img.split()[-1] if pil_img.mode == 'RGBA' else None)
                        pil_img = background
                    # Comprimir para JPEG em buffer de memória
                    buffer = BytesIO()
                    pil_img.save(buffer, format='JPEG', quality=85, optimize=True)
                    buffer.seek(0)
                    # Usar ImageReader do ReportLab (já importado como RLImage)
                    canvas.drawImage(RLImage(buffer), 0, 0, width=w, height=h,
                                     preserveAspectRatio=False, mask='auto')
                except Exception as e:
                    # Fallback para imagem original em caso de erro
                    print(f"  [AVISO] Falha na compressão da capa: {e}")
                    canvas.drawImage(img_capa, 0, 0, width=w, height=h,
                                     preserveAspectRatio=False, mask='auto')
            else:
                # Se Pillow não disponível, usar imagem original
                canvas.drawImage(img_capa, 0, 0, width=w, height=h,
                                 preserveAspectRatio=False, mask='auto')
        else:
            # Fallback: fundo azul caso a imagem não seja encontrada
            canvas.setFillColor(COR_PRIM)
            canvas.rect(0, 0, w, h, fill=1, stroke=0)
            canvas.setFillColor(COR_VERDE)
            canvas.rect(0, h * 0.48, w, 3*mm, fill=1, stroke=0)

        # ── Nome do Projeto — sobreposto sobre a imagem ──
        # Título Principal
        titulo_capa = f"Relatório Analítico de Apontamentos do {nome_projeto}"
        
        # Sombra sutil para legibilidade
        canvas.setFillColor(colors.HexColor("#00000055"))
        canvas.setFont("Helvetica-Bold", 20)
        canvas.drawCentredString(w/2 + 1, h * 0.55 - 1, titulo_capa)
        # Texto branco principal
        canvas.setFillColor(BRANCO)
        canvas.drawCentredString(w/2, h * 0.55, titulo_capa)

        # ── Data de Geração ──
        canvas.setFont("Helvetica", 14)
        canvas.setFillColor(colors.HexColor("#00000055"))
        canvas.drawCentredString(w/2 + 1, h * 0.50 - 1, f"Gerado em: {data_impressao}")
        canvas.setFillColor(BRANCO)
        canvas.drawCentredString(w/2, h * 0.50, f"Gerado em: {data_impressao}")

        # ── Rodapé Direito da Capa ──
        canvas.setFont("Helvetica", 7)
        canvas.drawRightString(w - 15*mm, 10*mm, "Desenvolvimento: Felipe de P. Luz")

        canvas.restoreState()
    return fn

def _on_page(nome_projeto, data_impressao):
    """Callback para páginas internas — com cabeçalho e rodapé."""
    def fn(canvas, doc):
        w, h = A4
        canvas.saveState()
        canvas.setFillColor(COR_PRIM)
        canvas.rect(0, h - 18*mm, w, 18*mm, fill=1, stroke=0)
        canvas.setFillColor(BRANCO)
        canvas.setFont("Helvetica-Bold", 11)
        canvas.drawString(15*mm, h - 11*mm, f"Relatório Analítico — {nome_projeto}")
        canvas.setFont("Helvetica", 8)
        canvas.setFillColor(COR_CINZA_M)
        canvas.drawRightString(w - 15*mm, h - 10*mm, f"Impresso em: {data_impressao}")
        canvas.drawRightString(w - 15*mm, h - 15*mm, f"Página {doc.page}")
        canvas.setStrokeColor(COR_VERDE)
        canvas.setLineWidth(1.5)
        canvas.line(0, h - 18*mm, w, h - 18*mm)
        # Rodapé
        canvas.setStrokeColor(COR_CINZA_M)
        canvas.setLineWidth(0.5)
        canvas.line(15*mm, 10*mm, w - 15*mm, 10*mm)
        canvas.setFont("Helvetica", 7)
        canvas.setFillColor(COR_CINZA_M)
        canvas.drawCentredString(w/2, 6*mm, f"{nome_projeto}")
        canvas.restoreState()
    return fn

# ─── Otimização de PDF ──────────────────────────────────────────────────────
def otimizar_pdf_final(caminho_entrada):
    """
    Tenta otimizar o PDF gerado reescrevendo-o com PyPDF2 (se disponível).
    A compressão de conteúdo (zlib) já é feita pelo ReportLab com pageCompression=1,
    mas o PyPDF2 pode ajudar a limpar metadados ou reestruturar o arquivo.
    """
    try:
        if 'PdfReader' not in globals() or 'PdfWriter' not in globals():
            return

        # Verificar tamanho original
        tamanho_orig = os.path.getsize(caminho_entrada)
        
        reader = PdfReader(caminho_entrada)
        writer = PdfWriter()

        # Copiar páginas (compressão de streams é padrão no PyPDF2 moderno)
        for page in reader.pages:
            writer.add_page(page)
        
        # Copiar metadados
        if reader.metadata:
            writer.add_metadata(reader.metadata)
        
        # Salvar em arquivo temporário
        caminho_temp = caminho_entrada.replace(".pdf", "_temp.pdf")
        with open(caminho_temp, "wb") as f:
            writer.write(f)
        
        tamanho_novo = os.path.getsize(caminho_temp)
        
        # Só substituir se houve redução ou se a diferença for desprezível (limpeza)
        # Se aumentou muito (as vezes acontece), descarta.
        if tamanho_novo < tamanho_orig:
            os.replace(caminho_temp, caminho_entrada)
            reducao = (1 - tamanho_novo/tamanho_orig) * 100
            print(f"  [INFO] Otimização PyPDF2: {tamanho_orig/1024:.1f}KB -> {tamanho_novo/1024:.1f}KB (-{reducao:.1f}%)")
        else:
            os.remove(caminho_temp)
            # print(f"  [INFO] Otimização PyPDF2 não reduziu tamanho ({tamanho_novo/1024:.1f}KB). Mantendo original.")
            
    except Exception as e:
        print(f"  [AVISO] Falha na otimização pós-processamento: {e}")
        if os.path.exists(caminho_entrada.replace(".pdf", "_temp.pdf")):
            try:
                os.remove(caminho_entrada.replace(".pdf", "_temp.pdf"))
            except:
                pass

def gerar_relatorio_analitico(arquivo_jsonl, pasta_base, data_str, data_impressao):
    nome_proj = os.path.splitext(os.path.basename(arquivo_jsonl))[0]
    pasta_out = os.path.join(pasta_base, nome_proj, data_str)
    os.makedirs(pasta_out, exist_ok=True)

    # Carregar dados
    dados = []
    with open(arquivo_jsonl, 'r', encoding='utf-8') as f:
        for l in f:
            if l.strip():
                try:
                    item = json.loads(l)
                    item['FixStatus_Ref'] = item.get('FixStatus') or item.get('Status', 'Open')
                    dados.append(item)
                except: pass

    if not dados:
        print(f"  [{nome_proj}] Sem dados. Pulando.")
        return

    df = pd.DataFrame(dados)
    analise, kpis = gerar_analise(df, nome_proj)
    st = estilos()

    caminho_pdf = os.path.join(pasta_out, f"{nome_proj}_Relatorio_Analitico.pdf")
    # Otimização: pageCompression=1 ativa compressão zlib dos fluxos de conteúdo
    doc = SimpleDocTemplate(caminho_pdf, pagesize=A4,
                              leftMargin=15*mm, rightMargin=15*mm,
                              topMargin=23*mm, bottomMargin=15*mm,
                              pageCompression=1)

    story = []
    largura = A4[0] - 30*mm

    # ── Página de Capa (conteúdo vazio — desenhada via canvas callback) ──
    story.append(Spacer(1, 1))  # placeholder para a capa
    story.append(PageBreak())

    # ── Índice ──
    st_indice_titulo = ParagraphStyle('idx_titulo', fontName='Helvetica-Bold', fontSize=11,
                                       textColor=COR_PRIM, leading=18, spaceBefore=2*mm)
    st_indice_item = ParagraphStyle('idx_item', fontName='Helvetica', fontSize=10,
                                     textColor=colors.HexColor('#333333'), leading=16,
                                     leftIndent=10*mm)

    story.append(Paragraph("Índice", st['titulo']))
    story.append(HRFlowable(width="100%", thickness=1, color=COR_VERDE, spaceAfter=4*mm))

    indice_items = [
        "1. Objetivo",
        "2. Resumo Executivo",
        "3. Segmentação: Segurança vs Qualidade",
        "4. Análise por Categoria e Subcategorias",
        "5. Tendência Temporal",
        "6. SLA / Tempo de Resposta (MTTR)",
        "7. Análise de Causa Raiz",
        "8. Conclusão",
        "9. Anexo — Relatório SnagR",
    ]
    for item in indice_items:
        story.append(Paragraph(item, st_indice_item))
    story.append(PageBreak())

    # ── Objetivo ──
    story.append(Paragraph("1. Objetivo", st['subtitulo']))
    story.append(Paragraph(
        f"Este relatório tem como objetivo apresentar uma análise detalhada dos "
        f"apontamentos de qualidade e segurança registrados no projeto <b>{nome_proj}</b>, "
        f"identificados através da plataforma SnagR."
        f"<br/><br/>"
        f"O documento abrange o diagnóstico do cenário atual, a avaliação de desempenho "
        f"por categoria, subcategoria e empreiteira, a análise de tendência temporal, "
        f"o tempo médio de resposta (SLA) e a identificação de padrões de causa raiz."
        f"<br/><br/>"
        f"As informações aqui consolidadas visam apoiar a tomada de decisão da gestão, "
        f"direcionando ações corretivas e preventivas para a melhoria contínua dos "
        f"processos construtivos e da segurança do trabalho.",
        st['corpo']
    ))

    st_termos_titulo = ParagraphStyle(
        'termos_titulo',
        fontName='Helvetica-Bold',
        fontSize=11,
        textColor=COR_PRIM,
        spaceBefore=4*mm,
        spaceAfter=2*mm
    )

    termos_defs = (
        "<b>Apontamento:</b> registro de uma não conformidade, risco ou oportunidade de melhoria identificado em campo.<br/>"
        "<b>KPI (Indicador-chave):</b> métrica objetiva usada para acompanhar desempenho (ex.: total de itens, pendências, taxa de resolução).<br/>"
        "<b>Backlog:</b> volume acumulado de apontamentos ainda não encerrados (pendentes).<br/>"
        "<b>Pendente:</b> apontamento sem tratativa final (ex.: “Sem Correção”).<br/>"
        "<b>Resolvido:</b> apontamento tratado e encerrado no fluxo (ex.: “Corrigido” e/ou “Verificado”).<br/>"
        "<b>Taxa de resolução:</b> percentual de apontamentos resolvidos em relação ao total no período analisado.<br/>"
        "<b>SLA (Tempo de resposta):</b> tempo observado entre a abertura e o encerramento/validação do apontamento, usado para medir agilidade de tratativa.<br/>"
        "<b>MTTR:</b> tempo médio para resolução; indicador de eficiência na correção (quanto menor, melhor).<br/>"
        "<b>Categoria / Subcategoria:</b> classificação do tema do apontamento (Categoria) e o tipo específico recorrente (Subcategoria/Defect).<br/>"
        "<b>Segmentação (Qualidade vs Segurança):</b> agrupamento dos apontamentos por natureza do impacto (integridade do produto vs risco de SST).<br/>"
        "<b>Causa raiz:</b> motivo estrutural que explica a recorrência do problema; base para ações preventivas e redução de reincidência."
    )

    story.append(KeepTogether([
        Paragraph("1.1 Termos e Definições", st_termos_titulo),
        Paragraph(termos_defs, st['corpo']),
    ]))
    story.append(Spacer(1, 4*mm))

    # ── KPIs em Cards ──
    def kpi_cell(valor, label, cor=COR_PRIM):
        return [
            Paragraph(str(valor), ParagraphStyle('kv', fontName='Helvetica-Bold', fontSize=22, textColor=cor, alignment=TA_CENTER)),
            Paragraph(label, st['kpi_label']),
        ]

    kpi_data = [
        [kpi_cell(kpis['total'], "TOTAL")[0], kpi_cell(kpis['resolvidos'], "RESOLVIDOS", colors.HexColor("#00c853"))[0],
         kpi_cell(kpis['sem_corr'], "PENDENTES", COR_VERM)[0], kpi_cell(f"{kpis['pct_res']:.0f}%", "RESOLUÇÃO", COR_VERDE)[0]],
        [kpi_cell(kpis['total'], "")[1], kpi_cell(kpis['resolvidos'], "")[1],
         kpi_cell(kpis['sem_corr'], "")[1], kpi_cell(f"{kpis['pct_res']:.0f}%", "")[1]],
    ]
    # Fix: rebuild as proper 2-row table
    kpi_data = [
        [Paragraph(str(kpis['total']), ParagraphStyle('k1', fontName='Helvetica-Bold', fontSize=24, textColor=COR_PRIM, alignment=TA_CENTER)),
         Paragraph(str(kpis['resolvidos']), ParagraphStyle('k2', fontName='Helvetica-Bold', fontSize=24, textColor=colors.HexColor("#00c853"), alignment=TA_CENTER)),
         Paragraph(str(kpis['sem_corr']), ParagraphStyle('k3', fontName='Helvetica-Bold', fontSize=24, textColor=COR_VERM, alignment=TA_CENTER)),
         Paragraph(f"{kpis['pct_res']:.0f}%", ParagraphStyle('k4', fontName='Helvetica-Bold', fontSize=24, textColor=COR_VERDE, alignment=TA_CENTER))],
        [Paragraph("Total", st['kpi_label']),
         Paragraph("Resolvidos", st['kpi_label']),
         Paragraph("Pendentes", st['kpi_label']),
         Paragraph("Taxa Resolução", st['kpi_label'])],
    ]
    quarter = largura / 4
    kpi_table = Table(kpi_data, colWidths=[quarter]*4, rowHeights=[12*mm, 6*mm])
    kpi_table.setStyle(TableStyle([
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ('ALIGN',  (0,0), (-1,-1), 'CENTER'),
        ('BOX',    (0,0), (-1,-1), 0.5, COR_CINZA_M),
        ('INNERGRID', (0,0), (-1,-1), 0.5, COR_CINZA_L),
        ('BACKGROUND', (0,0), (-1,0), COR_CINZA_L),
        ('TOPPADDING', (0,0), (-1,-1), 4),
        ('BOTTOMPADDING', (0,0), (-1,-1), 4),
    ]))

    # Seções de Análise
    secoes_texto = [
        ("2. Resumo Executivo",                    analise['resumo']),
        ("3. Segmentação: Segurança vs Qualidade", analise['segmentacao']),
        ("5. Tendência Temporal",                  analise['tendencia']),
        ("6. SLA / Tempo de Resposta (MTTR)",      analise['sla']),
        ("7. Análise de Causa Raiz",               analise['causa_raiz']),
    ]

    # Estilo de tabela genérico
    def get_table_style():
        return TableStyle([
            ('BACKGROUND', (0,0), (-1,0), COR_PRIM),
            ('TEXTCOLOR',  (0,0), (-1,0), BRANCO),
            ('FONTNAME',   (0,0), (-1,0), 'Helvetica-Bold'),
            ('FONTSIZE',   (0,0), (-1,0), 9),
            ('ALIGN',      (0,0), (-1,0), 'CENTER'),
            ('VALIGN',     (0,0), (-1,-1), 'MIDDLE'),
            ('GRID',       (0,0), (-1,-1), 0.5, COR_CINZA_M),
            ('ROWBACKGROUNDS', (0,1), (-1,-1), [BRANCO, COR_CINZA_L]),
            ('FONTSIZE',   (0,1), (-1,-1), 8),
        ])

    # Seção 2 — Resumo Executivo
    story.append(KeepTogether([
        Paragraph(secoes_texto[0][0], st['subtitulo']),
        Paragraph(secoes_texto[0][1], st['corpo']),
        Spacer(1, 2*mm)
    ]))

    # KPIs ao final do texto do Resumo Executivo
    story.append(kpi_table)
    story.append(Spacer(1, 6*mm))

    # Seção 3 — Segmentação (antes do gráfico de categorias)
    story.append(KeepTogether([
        Paragraph(secoes_texto[1][0], st['subtitulo']),
        Paragraph(secoes_texto[1][1], st['corpo']),
        Spacer(1, 2*mm)
    ]))

    # ── Seção 4 — Mini Gráficos de Barras por Categoria/Subcategoria ──
    story.append(Paragraph("4. Análise por Categoria e Subcategorias", st['subtitulo']))
    story.append(Paragraph(analise['cat_intro'], st['corpo']))
    story.append(Spacer(1, 3*mm))
    for fl in montar_secao_categorias(df, largura):
        story.append(fl)
    story.append(Spacer(1, 4*mm))
    story.append(Paragraph(analise['cat_resumo'], st['corpo']))
    story.append(Spacer(1, 2*mm))

    # Seções 5 em diante (texto)
    for titulo_sec, conteudo in secoes_texto[2:]:
        # Especial para Causa Raiz: Introdução e Conclusão
        if "7. Análise de Causa Raiz" in titulo_sec:
            sec_story = [
                Paragraph(titulo_sec, st['subtitulo']),
                Paragraph(analise.get('causa_raiz_intro', ''), st['corpo']),
                Spacer(1, 4*mm),
                Paragraph(conteudo, st['corpo']),
                Spacer(1, 4*mm),
                # Paragraph(analise.get('causa_raiz_conclusao', ''), st['corpo']), # Removido
                Spacer(1, 2*mm)
            ]
        else:
            sec_story = [
                Paragraph(titulo_sec, st['subtitulo']),
                Paragraph(conteudo, st['corpo'])
            ]
        
        if "5. Tendência Temporal" in titulo_sec:
            for fl in montar_grafico_tendencia(df, largura):
                sec_story.append(fl)

        # Especial para SLA: Tabelas de MTTR e pendências antigas
        if "6. SLA" in titulo_sec:
            st_tbl_preto_sla = ParagraphStyle('tbl_preto_sla', fontName='Helvetica', fontSize=7,
                                              textColor=colors.HexColor("#000000"), alignment=TA_CENTER)
            # MTTR por Categoria
            if analise.get('sla_cat_data'):
                sec_story.append(Spacer(1, 4*mm))
                sec_story.append(Paragraph("<b>MTTR por Categoria (mais lentas):</b>", st['destaque']))
                t_cat = [["Categoria", "Média", "Mediana", "Qtd", "Status"]]
                for item in analise['sla_cat_data']:
                    st_txt = item['Status']
                    if st_txt == "Crítico":
                        st_txt = f"<font color='#ff5252'>{st_txt}</font>"
                    t_cat.append([
                        Paragraph(item['Categoria'], st_tbl_preto_sla),
                        item['Média'], item['Mediana'], str(item['Qtd']),
                        Paragraph(st_txt, st['rodape'])
                    ])
                tab_cat = Table(t_cat, colWidths=[largura*0.4, largura*0.15, largura*0.15, largura*0.1, largura*0.2])
                tab_cat.setStyle(get_table_style())
                sec_story.append(tab_cat)

            # MTTR por Grupo
            if analise.get('sla_grp_data'):
                sec_story.append(Spacer(1, 4*mm))
                sec_story.append(Paragraph("<b>MTTR por Grupo/Empreiteira (mais lentas):</b>", st['destaque']))
                t_grp = [["Grupo/Empreiteira", "Média", "Mediana", "Qtd", "Status"]]
                for item in analise['sla_grp_data']:
                    st_txt = item['Status']
                    if st_txt == "Crítico":
                        st_txt = f"<font color='#ff5252'>{st_txt}</font>"
                    t_grp.append([
                        Paragraph(item['Grupo'], st_tbl_preto_sla),
                        item['Média'], item['Mediana'], str(item['Qtd']),
                        Paragraph(st_txt, st['rodape'])
                    ])
                tab_grp = Table(t_grp, colWidths=[largura*0.4, largura*0.15, largura*0.15, largura*0.1, largura*0.2])
                tab_grp.setStyle(get_table_style())
                sec_story.append(tab_grp)

            # Tabela de pendências antigas
            if analise.get('top_antigos'):
                sec_story.append(Spacer(1, 4*mm))
                sec_story.append(Paragraph("<b>Os 20 apontamentos pendentes mais antigos:</b>", st['destaque']))
                
                st_tbl_preto = ParagraphStyle('tbl_preto', fontName='Helvetica', fontSize=7,
                                              textColor=colors.HexColor("#000000"), alignment=TA_CENTER)
                t_data = [["ID", "Dias", "Categoria", "Defeito", "Localização"]]
                for item in analise['top_antigos']:
                    t_data.append([
                        str(item['ID']),
                        str(item['Dias']),
                        Paragraph(item['Categoria'][:30], st_tbl_preto),
                        Paragraph(item['Defeito'][:30], st_tbl_preto),
                        Paragraph(item['Local'][:30], st_tbl_preto)
                    ])
                
                tab = Table(t_data, colWidths=[largura*0.1, largura*0.1, largura*0.25, largura*0.25, largura*0.3], repeatRows=1)
                tab.setStyle(get_table_style())
                sec_story.append(tab)

        sec_story.append(Spacer(1, 2*mm))
        story.append(KeepTogether(sec_story))

    # ── Conclusão Final ──
    story.append(PageBreak())
    story.append(Paragraph("8. Conclusão", st['subtitulo']))
    conclusao = gerar_conclusao(analise, kpis, nome_proj)
    # Remover o sufixo automático se o texto veio do LLM (identificado se não tiver a marcação de rodapé)
    # Mas como o LLM não gera a marcação, apenas usamos o texto.
    
    story.append(Paragraph(conclusao, st['corpo']))
    
    # Adicionar disclaimer discreto apenas se NÃO for LLM (opcional, mas o usuário pediu para não ter marca de IA, então não adicionamos nada extra)
    
    story.append(Spacer(1, 10*mm))
    story.append(HRFlowable(width="100%", thickness=1.5, color=COR_VERDE, spaceAfter=4*mm))
    story.append(Paragraph(
        f"<b>{nome_proj}</b> — Relatório gerado em {data_impressao}",
        ParagraphStyle('fim', fontName='Helvetica', fontSize=9,
                        textColor=COR_CINZA_M, alignment=TA_CENTER)
    ))

    doc.build(story, onFirstPage=_on_cover(nome_proj, data_impressao),
              onLaterPages=_on_page(nome_proj, data_impressao))
    
    # Otimização pós-geração (PyPDF2)
    otimizar_pdf_final(caminho_pdf)
    
    print(f"  [OK] {nome_proj}: {len(dados)} apontamentos -> {caminho_pdf}")

    # Gerar Word (.docx)
    if 'Document' in globals():
        gerar_docx_analitico(df, analise, kpis, nome_proj, pasta_out, data_impressao)

# ─── Gerar Arquivo Word (DOCX) ──────────────────────────────────────────────
def gerar_docx_analitico(df, analise, kpis, nome_proj, pasta_out, data_impressao):
    """Gera uma versão Word do relatório analítico."""
    try:
        doc = Document()
        
        # Título
        t = doc.add_heading(f"Relatório Analítico — {nome_proj}", 0)
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f"Gerado em: {data_impressao}").alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_page_break()
        
        # 1. Objetivo
        doc.add_heading("1. Objetivo", level=1)
        doc.add_paragraph(
            f"Este relatório tem como objetivo apresentar uma análise detalhada dos "
            f"apontamentos de qualidade e segurança registrados no projeto {nome_proj}, "
            f"identificados através da plataforma SnagR."
        )
        
        # Resumo de KPIs
        doc.add_heading("KPIs Globais", level=2)
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'TOTAL'
        hdr_cells[1].text = 'RESOLVIDOS'
        hdr_cells[2].text = 'PENDENTES'
        hdr_cells[3].text = 'RESOLUÇÃO'
        
        row_cells = table.add_row().cells
        row_cells[0].text = str(kpis['total'])
        row_cells[1].text = str(kpis['resolvidos'])
        row_cells[2].text = str(kpis['sem_corr'])
        row_cells[3].text = f"{kpis['pct_res']:.1f}%"
        
        # Seções de texto
        secoes = [
            ("2. Resumo Executivo", analise['resumo']),
            ("3. Segmentação: Segurança vs Qualidade", analise['segmentacao']),
            ("4. Análise por Categoria e Subcategorias", "Ver detalhes no PDF para visualização gráfica."),
            ("5. Tendência Temporal", analise['tendencia']),
            ("6. SLA / Tempo de Resposta (MTTR)", analise['sla']),
            ("7. Análise de Causa Raiz", analise['causa_raiz']),
            ("8. Conclusão", gerar_conclusao(analise, kpis, nome_proj))
        ]
        
        for titulo, conteudo in secoes:
            doc.add_heading(titulo, level=1)
            # Limpar tags HTML básicas que podem estar no texto
            txt = conteudo.replace("<b>", "").replace("</b>", "").replace("<br/>", "\n").replace("<i>", "").replace("</i>", "").replace("<font color='#ff5252'>", "").replace("</font>", "")
            doc.add_paragraph(txt)
            
            # Adicionar tabelas de SLA se for a seção 6
            if "6. SLA" in titulo:
                if analise.get('sla_cat_data'):
                    doc.add_heading("MTTR por Categoria", level=2)
                    t_cat = doc.add_table(rows=1, cols=5)
                    t_cat.style = 'Table Grid'
                    h = t_cat.rows[0].cells
                    h[0].text, h[1].text, h[2].text, h[3].text, h[4].text = "Categoria", "Média", "Mediana", "Qtd", "Status"
                    for item in analise['sla_cat_data']:
                        r = t_cat.add_row().cells
                        r[0].text, r[1].text, r[2].text, r[3].text, r[4].text = item['Categoria'], item['Média'], item['Mediana'], str(item['Qtd']), item['Status']
                
                if analise.get('top_antigos'):
                    doc.add_heading("Os 20 itens mais antigos", level=2)
                    t_ant = doc.add_table(rows=1, cols=5)
                    t_ant.style = 'Table Grid'
                    h = t_ant.rows[0].cells
                    h[0].text, h[1].text, h[2].text, h[3].text, h[4].text = "ID", "Dias", "Categoria", "Defeito", "Local"
                    for item in analise['top_antigos']:
                        r = t_ant.add_row().cells
                        r[0].text, r[1].text, r[2].text, r[3].text, r[4].text = str(item['ID']), str(item['Dias']), item['Categoria'], item['Defeito'], item['Local']

        caminho_docx = os.path.join(pasta_out, f"{nome_proj}_Relatorio_Analitico.docx")
        doc.save(caminho_docx)
        print(f"  [OK] Word gerado: {os.path.basename(caminho_docx)}")
    except Exception as e:
        print(f"  [ERRO] Falha ao gerar Word: {e}")

# ─── Main ────────────────────────────────────────────────────────────────────
def main():
    print("=" * 58)
    print("  GERADOR DE RELATORIO ANALITICO - SnagR (IA Nativa)")
    print("=" * 58)

    agora = datetime.now()
    data_str = agora.strftime("%Y-%m-%d_%H-%M")
    data_imp = agora.strftime("%d/%m/%Y %H:%M")

    pasta_base = os.path.join(os.getcwd(), PASTA_RELATORIOS)
    os.makedirs(pasta_base, exist_ok=True)

    arquivos = sorted(glob.glob("*.jsonl"))
    if not arquivos:
        print("ERRO: Nenhum .jsonl encontrado!")
        return

    print(f"  {len(arquivos)} projeto(s). Gerando relatórios analíticos...\n")
    for arq in arquivos:
        print(f"\n  [{os.path.splitext(os.path.basename(arq))[0]}] Iniciando...")
        gerar_relatorio_analitico(arq, pasta_base, data_str, data_imp)

    print(f"\n  Salvos em: {pasta_base}")
    print("=" * 58)

if __name__ == "__main__":
    try:
        main()
    except Exception as e:
        print(f"\nERRO FATAL: {e}")
        import traceback
        traceback.print_exc()
